import docx

doc1 = docx.Document('input.docx')
paragraphs = len(doc1.paragraphs)
tables = len(doc1.tables)
text = []
for i in doc1.paragraphs:
    if i.text != '':
        text.append(i.text)

slova = []
for i in range(len(text)):
    a = text[i].split(' ')
    for j in a:
        slova.append(j)
len_sliv = len(slova)
for table in doc1.tables:
    for raw in table.rows:
        for cell in raw.cells:
            if cell.text != '':
                text.append(cell.text)

simvolika = []
for i in text:
    for j in i:
        if j != '' and j != ' ':
            simvolika.append(j)
len_symv = len(simvolika)
print(f'Кількість символів : {len_symv}\n'
      f'Кількість слів : {len_sliv}\n'
      f'Кількість параграфів : {paragraphs}\n'
      f'Кількість таблиць : {tables}\n')
